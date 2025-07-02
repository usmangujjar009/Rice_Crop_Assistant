import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

# ----------------------------
# Crop Thresholds
# ----------------------------
rice_thresholds = {
    "temperature": (20, 38, "Apply mulch or reduce watering", "Install shade net"),
    "moisture": (60, 80, "Increase watering", "Improve soil drainage"),
    "humidity": (70, 90, "Use sprinkler system", "Increase ventilation"),
    "pH": (5.5, 7.0, "Add lime to raise pH", "Use sulfur or compost to lower pH")
}

# ----------------------------
# Tkinter GUI Setup
# ----------------------------
root = tk.Tk()
root.title("🌾 Smart Rice Crop Advisor")
root.state('zoomed')
root.configure(bg="#eefaf7")

# ----------------------------
# Title
# ----------------------------
tk.Label(root, text="🌾 Smart Rice Crop Advisor", font=("Helvetica", 30, "bold"),
         bg="#eefaf7", fg="#004d40").pack(pady=30)

# ----------------------------
# Input Fields
# ----------------------------
input_frame = tk.Frame(root, bg="#eefaf7")
input_frame.pack(pady=10)

def create_input(label, row):
    tk.Label(input_frame, text=label, font=("Arial", 14, "bold"), bg="#eefaf7")\
        .grid(row=row*2, column=0, sticky="w", padx=30, pady=5)
    entry = tk.Entry(input_frame, font=("Arial", 14), width=40)
    entry.grid(row=row*2+1, column=0, padx=30, pady=5)
    return entry

entry_temp = create_input("🌡️ Temperature (°C) [20–38]:", 0)
entry_moisture = create_input("💧 Moisture (%) [60–80]:", 1)
entry_humidity = create_input("💨 Humidity (%) [70–90]:", 2)
entry_ph = create_input("⚗️ pH Level [5.5–7.0]:", 3)

# ----------------------------
# Result Text Box
# ----------------------------
tk.Label(root, text="📋 Crop Condition Suggestions", font=("Arial", 20, "bold"),
         bg="#eefaf7", fg="#004d40").pack()

result_text = tk.Text(root, height=18, width=100, font=("Courier New", 14, "bold"),
                      wrap=tk.WORD, borderwidth=3, relief="groove")
result_text.pack(padx=30, pady=10)
result_text.config(state='disabled')
result_text.tag_config("red", foreground="#c62828")
result_text.tag_config("orange", foreground="#f57c00")
result_text.tag_config("green", foreground="#2e7d32")

# ----------------------------
# Check Crop Condition
# ----------------------------
def check_condition(param, value):
    min_val, max_val, low_suggestion, high_suggestion = rice_thresholds[param]
    if value < min_val:
        return f"⚠️ {param.capitalize()} is LOW ({value})\n💡 Suggestion: {low_suggestion}\n", "red"
    elif value > max_val:
        return f"⚠️ {param.capitalize()} is HIGH ({value})\n💡 Suggestion: {high_suggestion}\n", "orange"
    else:
        return f"✅ {param.capitalize()} is OK ({value})\n", "green"

# ----------------------------
# Analyze Button Function
# ----------------------------
def analyze():
    try:
        temp = float(entry_temp.get())
        moisture = float(entry_moisture.get())
        humidity = float(entry_humidity.get())
        ph = float(entry_ph.get())

        result_text.config(state='normal')
        result_text.delete('1.0', tk.END)

        for param, val in [("temperature", temp), ("moisture", moisture), ("humidity", humidity), ("pH", ph)]:
            result, color = check_condition(param, val)
            result_text.insert(tk.END, result + "\n", color)

        result_text.config(state='disabled')

    except ValueError:
        messagebox.showerror("Input Error", "❌ Please enter valid numeric values only!")

# ----------------------------
# Export to Word File
# ----------------------------
def export_to_docx():
    result_text.config(state='normal')
    content = result_text.get("1.0", tk.END).strip()
    result_text.config(state='disabled')

    if not content:
        messagebox.showwarning("Export Failed", "⚠️ No result to export!")
        return

    doc = Document()

    # 🖼️ Add Logo if present
    logo_path = "logo.png"
    if os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(2.5))
        except Exception as e:
            print("Logo error:", e)

    # 🧾 Title
    heading = doc.add_heading("🌾 Smart Rice Crop Report", 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    # 📋 Add each line with color
    for line in content.split('\n'):
        if not line.strip():
            continue

        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True

        if "LOW" in line:
            run.font.color.rgb = RGBColor(198, 40, 40)   # Red
        elif "HIGH" in line:
            run.font.color.rgb = RGBColor(245, 124, 0)   # Orange
        elif "OK" in line:
            run.font.color.rgb = RGBColor(46, 125, 50)   # Green
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)       # Black

    # 📅 Footer
    doc.add_paragraph("\n")
    footer = doc.add_paragraph(f"🕒 Report generated on: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}")
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # 💾 Save
    doc.save("rice_report_2025.docx")
    messagebox.showinfo("Export Successful", "✅ Report saved as 'rice_report_2025.docx'!")

# ----------------------------
# Buttons
# ----------------------------
button_frame = tk.Frame(root, bg="#eefaf7")
button_frame.pack(pady=20)

tk.Button(button_frame, text="📊 Analyze Crop Conditions", font=("Arial", 16, "bold"),
          bg="#007f5f", fg="white", padx=30, pady=12, command=analyze)\
    .grid(row=0, column=0, padx=10)

tk.Button(button_frame, text="📝 Export Report", font=("Arial", 16, "bold"),
          bg="#0080ff", fg="white", padx=30, pady=12, command=export_to_docx)\
    .grid(row=0, column=1, padx=10)

tk.Button(button_frame, text="❌ Exit", font=("Arial", 16, "bold"),
          bg="#ff4d4d", fg="white", padx=30, pady=12, command=root.destroy)\
    .grid(row=0, column=2, padx=10)

# ----------------------------
# Start GUI
# ----------------------------
root.mainloop()

